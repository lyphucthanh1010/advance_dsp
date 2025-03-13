# Nếu chưa cài đặt python-docx, hãy cài đặt bằng:
# pip install python-docx

from docx import Document

# Tạo đối tượng Document
document = Document()

# Thêm tiêu đề chính
document.add_heading('Hướng dẫn Workflow tự động thu thập dữ liệu từ News API và lưu vào Google Sheets trong n8n', 0)

# Phần 1: Giới thiệu
document.add_heading('1. Giới thiệu', level=1)
document.add_paragraph(
    'Workflow này được xây dựng trên n8n nhằm tự động hóa quá trình thu thập dữ liệu từ News API và lưu trữ vào Google Sheets. '
    'Mỗi ngày vào lúc 9h sáng, workflow sẽ kích hoạt, gửi request lấy dữ liệu các bản tin của ngày hôm qua, xử lý dữ liệu và ghi vào Google Sheets theo định dạng mong muốn.'
)

# Phần 2: Các bước chi tiết trong Workflow
document.add_heading('2. Các bước chi tiết trong Workflow', level=1)

# Bước 1: Node Cron
document.add_heading('Bước 1: Node Cron', level=2)
document.add_paragraph(
    'Mục đích: Lập lịch chạy workflow tự động vào lúc 9h sáng hàng ngày.\n'
    'Cấu hình:\n'
    '- Thêm node Cron vào workflow.\n'
    '- Chọn chế độ "Every Day".\n'
    '- Đặt Hour = 9 và Minute = 0.\n'
    '- Lưu lại cấu hình.'
)

# Bước 2: Node HTTP Request
document.add_heading('Bước 2: Node HTTP Request', level=2)
document.add_paragraph(
    'Mục đích: Gửi yêu cầu GET đến News API để lấy dữ liệu bản tin.\n'
    'Cấu hình:\n'
    '- Thêm node HTTP Request sau node Cron.\n'
    '- Method: GET\n'
    '- URL: https://newsapi.org/v2/everything\n'
    '- Query Parameters:\n'
    '    • apiKey: API key của bạn\n'
    '    • from: Sử dụng expression {{$moment().subtract(1, \'days\').format("YYYY-MM-DD")}} để lấy ngày hôm qua\n'
    '    • to: Sử dụng expression tương tự để đặt ngày kết thúc\n'
    '    • Các tham số khác (q, language, …) tùy nhu cầu\n'
    '- Cấu hình thêm Header nếu API yêu cầu.'
)

# Bước 3: Node Function
document.add_heading('Bước 3: Node Function', level=2)
document.add_paragraph(
    'Mục đích: Xử lý dữ liệu JSON trả về từ News API, lọc và định dạng các thông tin cần thiết.\n'
    'Ví dụ mã nguồn trong Node Function:\n'
    '--------------------------------------------------\n'
    '// Giả sử response từ HTTP Request có cấu trúc: { articles: [ { title, description, publishedAt, source: { name } }, ... ] }\n'
    'const articles = items[0].json.articles || [];\n'
    'return articles.map(article => {\n'
    '  return {\n'
    '    json: {\n'
    '      title: article.title,\n'
    '      description: article.description,\n'
    '      publishedAt: article.publishedAt,\n'
    '      source: article.source ? article.source.name : ""\n'
    '    }\n'
    '  };\n'
    '});\n'
    '--------------------------------------------------'
)

# Bước 4: Node Google Sheets
document.add_heading('Bước 4: Node Google Sheets', level=2)
document.add_paragraph(
    'Mục đích: Ghi dữ liệu đã xử lý vào bảng tính Google Sheets.\n'
    'Cấu hình:\n'
    '- Thêm node Google Sheets sau node Function.\n'
    '- Xác thực với Google API qua OAuth2, đảm bảo quyền truy cập vào Google Sheets.\n'
    '- Chọn Spreadsheet ID của bảng tính cần ghi dữ liệu.\n'
    '- Chọn Worksheet Name (ví dụ: "NewsData").\n'
    '- Operation: Chọn "Append" để thêm dữ liệu vào cuối bảng.\n'
    '- Mapping: Gán các trường dữ liệu (title, description, publishedAt, source) từ node Function với các cột tương ứng.'
)

# Bước 5: Kết nối các Node
document.add_heading('Bước 5: Kết nối các Node', level=2)
document.add_paragraph(
    'Sắp xếp các node theo trình tự:\n'
    'Cron → HTTP Request → Function → Google Sheets\n'
    'Kiểm tra kết nối giữa các node để đảm bảo dữ liệu được truyền đúng từ bước này sang bước khác.'
)

# Phần 3: Sơ đồ Workflow
document.add_heading('3. Sơ đồ Workflow', level=1)
document.add_paragraph(
    'Sơ đồ dưới đây mô tả luồng dữ liệu:\n'
    '[ Cron (Trigger: 9h sáng hàng ngày) ]\n'
    '         │\n'
    '         ▼\n'
    '[ HTTP Request (GET: News API) ]\n'
    '         │\n'
    '         ▼\n'
    '[ Function (Xử lý dữ liệu: Lọc và định dạng bản tin) ]\n'
    '         │\n'
    '         ▼\n'
    '[ Google Sheets (Append dữ liệu vào bảng tính) ]'
)

# Phần 4: Kiểm tra và Debug
document.add_heading('4. Kiểm tra và Debug', level=1)
document.add_paragraph(
    'Sau khi cấu hình xong workflow:\n'
    '- Sử dụng chức năng "Execute Workflow" của n8n để chạy thử và kiểm tra đầu ra của từng node.\n'
    '- Kiểm tra log từng node để xác định lỗi và đảm bảo dữ liệu được truyền đúng.\n'
    '- Điều chỉnh các tham số hoặc code trong Node Function nếu cần.'
)

# Phần 5: Các lưu ý khác
document.add_heading('5. Các lưu ý khác', level=1)
document.add_paragraph(
    '- Bảo mật API key của News API.\n'
    '- Kiểm tra định dạng ngày (from, to) phù hợp với yêu cầu của API.\n'
    '- Đảm bảo kết nối OAuth2 với Google Sheets hoạt động ổn định.'
)

# Phần 6: Kết luận
document.add_heading('6. Kết luận', level=1)
document.add_paragraph(
    'Workflow này tự động thu thập dữ liệu bản tin từ News API và ghi vào Google Sheets, giảm thiểu công việc thủ công và đảm bảo dữ liệu được cập nhật hàng ngày một cách chính xác và hiệu quả.'
)

document.save("workflow_n8n_detail.docx")
